"""
家装产教融合AI闭环工作流 - Flask Web应用（测试版本）
"""
from flask import Flask, render_template, request, jsonify
import os
import chardet
import datetime

# 创建Flask应用
app = Flask(__name__, template_folder='templates')

# 配置
app.config['UPLOAD_FOLDER'] = 'temp'
app.config['SUBMISSIONS_FOLDER'] = 'submissions'

# 创建必要的目录
os.makedirs('temp', exist_ok=True)
os.makedirs('submissions', exist_ok=True)

# 工作流状态存储
workflow_state = {
    'stage': 2,  # 测试版本：设置为2，允许用户直接进入学生端
    'project_material': None,
    'course_content': None,
    'student_tasks': [],
    'submissions': [],
    'feedback_report': None
}

def read_file_with_encoding(file_path):
    """
    读取文件内容，自动检测编码
    支持多种编码：UTF-8, GBK, GB2312, BIG5 等
    支持 Word 文档（.docx）内容提取
    注意：纯二进制文件（图片、PDF等）将不读取内容
    """
    import mimetypes
    
    # 获取文件扩展名
    filename = os.path.basename(file_path)
    file_ext = os.path.splitext(filename)[1].lower()
    
    # 定义 Word 文档扩展名
    word_extensions = {'.docx'}
    
    # 定义纯二进制文件扩展名列表（不解析）
    binary_extensions = {
        # 图片
        '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.ico', '.svg', '.webp', '.tiff',
        # 文档（纯二进制，不解析）
        '.pdf', '.doc', '.xls', '.xlsx', '.ppt', '.pptx',
        # 压缩文件
        '.zip', '.rar', '.7z', '.tar', '.gz',
        # 视频
        '.mp4', '.avi', '.mov', '.wmv', '.flv', '.mkv',
        # 音频
        '.mp3', '.wav', '.flac', '.aac', '.ogg',
        # 其他
        '.exe', '.dll', '.so', '.dylib'
    }
    
    # 处理 Word 文档
    if file_ext in word_extensions:
        try:
            print(f"[DEBUG] 开始解析Word文档: {filename}")
            from docx import Document
            doc = Document(file_path)
            
            # 提取所有段落文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            print(f"[DEBUG] 提取到 {len(paragraphs)} 个段落")
            
            # 提取表格内容
            tables = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables.append(' | '.join(row_text))
            
            print(f"[DEBUG] 提取到 {len(tables)} 个表格行")
            
            # 组合内容
            content = '\n\n'.join(paragraphs)
            if tables:
                content += '\n\n【表格内容】\n' + '\n'.join(tables)
            
            print(f"[DEBUG] Word文档内容长度: {len(content)} 字符")
            
            if content.strip():
                return content
            else:
                return f"[Word文档为空: {filename}]"
        except Exception as e:
            print(f"[DEBUG] 读取Word文档失败: {str(e)}")
            return f"[读取Word文档失败: {filename}，错误: {str(e)}]"
    
    # 检查是否是纯二进制文件
    if file_ext in binary_extensions:
        return f"[二进制文件，已成功上传: {filename}]"
    
    # 尝试读取文本文件
    try:
        # 先尝试 UTF-8
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            # 验证读取的内容是否包含过多控制字符（可能是二进制文件）
            control_chars = sum(1 for c in content if ord(c) < 32 and c not in '\n\r\t')
            if len(content) > 0 and control_chars / len(content) > 0.1:
                return f"[二进制文件，已成功上传: {filename}]"
            return content
    except UnicodeDecodeError:
        try:
            # 检测文件编码
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding']
                confidence = result['confidence']
                
                # 使用检测到的编码读取文件
                if encoding and confidence > 0.5:
                    decoded_content = raw_data.decode(encoding, errors='ignore')
                    # 验证读取的内容是否包含过多控制字符
                    control_chars = sum(1 for c in decoded_content if ord(c) < 32 and c not in '\n\r\t')
                    if len(decoded_content) > 0 and control_chars / len(decoded_content) > 0.1:
                        return f"[可能是二进制文件，已成功上传: {filename}]"
                    return decoded_content
                else:
                    # 尝试常见的中文编码
                    for enc in ['gbk', 'gb2312', 'gb18030', 'big5']:
                        try:
                            decoded_content = raw_data.decode(enc, errors='ignore')
                            control_chars = sum(1 for c in decoded_content if ord(c) < 32 and c not in '\n\r\t')
                            if len(decoded_content) > 0 and control_chars / len(decoded_content) <= 0.1:
                                return decoded_content
                        except:
                            continue
                    
                    # 如果都失败，返回文件信息
                    file_size = len(raw_data)
                    return f"[文件已成功上传: {filename}，大小: {file_size} 字节]"
        except Exception as e:
            # 如果所有编码都失败，返回文件信息
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                file_size = len(raw_data)
                return f"[文件已成功上传: {filename}，大小: {file_size} 字节]"
    except Exception as e:
        return f"[读取文件失败: {str(e)}]"

@app.route('/')
def index():
    """首页 - 三阶段工作流界面"""
    return render_template('index.html')

@app.route('/get_state')
def get_state():
    """获取工作流状态"""
    return jsonify(workflow_state)

@app.route('/upload_project', methods=['POST'])
def upload_project():
    """企业端：上传项目（测试版本）"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '没有选择文件'}), 400
    
    # 保存文件
    filename = file.filename
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)
    
    # 读取文件内容（使用编码检测）
    try:
        # 获取文件扩展名
        file_ext = os.path.splitext(filename)[1].lower()
        
        file_content = read_file_with_encoding(file_path)
        
        # 获取文件信息
        file_size = os.path.getsize(file_path)
        
        # 更新工作流状态（只存储文件信息，不存储内容）
        workflow_state['project_material'] = {
            'filename': filename,
            'file_size': file_size,
            'file_path': file_path
        }
        workflow_state['stage'] = 2
        
        # 测试版本：不调用工作流，直接返回模拟数据
        # 构建显示内容
        if file_content.startswith('[') and file_content.endswith(']'):
            # 二进制文件或特殊格式，只显示文件信息
            display_content = f'✅ 文件上传成功！\n\n文件名: {filename}\n文件大小: {file_size} 字节\n文件类型: 二进制文件（图片、PDF等）\n\n{file_content}'
        else:
            # 文本文件或Word文档，显示内容预览
            preview = file_content[:2000] + '...' if len(file_content) > 2000 else file_content
            file_type = "Word文档" if file_ext == '.docx' else "文本文件"
            display_content = f'✅ 文件读取成功！\n\n文件名: {filename}\n文件大小: {file_size} 字节\n文件类型: {file_type}\n\n文件内容预览:\n{preview}'
        
        return jsonify({
            'success': True,
            'data': {
                'course_content': display_content,
                'student_tasks': [
                    {'任务ID': '1', '任务名称': '测试任务1', '任务描述': '测试描述1', '难度': '简单', '截止时间': '2024-12-31'},
                    {'任务ID': '2', '任务名称': '测试任务2', '任务描述': '测试描述2', '难度': '中等', '截止时间': '2024-12-31'}
                ]
            }
        })
    except Exception as e:
        return jsonify({'error': f'处理文件失败: {str(e)}'}), 500

@app.route('/submit_homework', methods=['POST'])
def submit_homework():
    """学生端：提交作业（测试版本）"""
    print(f"[DEBUG] submit_homework 被调用")
    print(f"[DEBUG] request.files: {request.files}")
    print(f"[DEBUG] request.form: {request.form}")
    
    if 'file' not in request.files:
        print(f"[DEBUG] 没有找到 file 字段")
        return jsonify({'error': '没有上传文件'}), 400
    
    file = request.files['file']
    print(f"[DEBUG] 文件对象: {file}")
    print(f"[DEBUG] 文件名: {file.filename}")
    
    if file.filename == '':
        print(f"[DEBUG] 文件名为空")
        return jsonify({'error': '没有选择文件'}), 400
    
    # 保存作业文件
    filename = file.filename
    file_path = os.path.join(app.config['SUBMISSIONS_FOLDER'], filename)
    print(f"[DEBUG] 准备保存文件到: {file_path}")
    
    try:
        file.save(file_path)
        print(f"[DEBUG] 文件保存成功")
    except Exception as e:
        print(f"[DEBUG] 文件保存失败: {str(e)}")
        return jsonify({'error': f'文件保存失败: {str(e)}'}), 500
    
    # 预设的评价维度和修改建议库
    evaluation_dimensions = {
        '功能布局': {
            'score': 0,
            'threshold': 6,
            'suggestions': {
                'direction': '优化空间动线和功能分区',
                'cases': [
                    {'type': 'text', 'content': '参考案例：客厅-餐厅-厨房一字型布局，缩短动线，提升使用效率'},
                    {'type': 'text', 'content': '参考案例：主卧设置独立衣帽间，提升收纳功能'},
                    {'type': 'text', 'content': '参考案例：卫生间干湿分离设计，提升使用舒适度'}
                ]
            }
        },
        '风格融合': {
            'score': 0,
            'threshold': 6,
            'suggestions': {
                'direction': '增强设计风格与业主偏好的融合',
                'cases': [
                    {'type': 'text', 'content': '参考案例：新中式风格融入红松纹理与现代留白设计'},
                    {'type': 'text', 'content': '参考案例：北欧风格结合原木元素与莫兰迪色系'},
                    {'type': 'text', 'content': '参考案例：现代简约风格运用黑灰白对比色块'}
                ]
            }
        },
        '文化落地': {
            'score': 0,
            'threshold': 6,
            'suggestions': {
                'direction': '强化地域文化与现代生活的结合',
                'cases': [
                    {'type': 'text', 'content': '参考案例：东北火炕元素与现代地暖系统结合'},
                    {'type': 'text', 'content': '参考案例：江南园林造景手法融入室内空间'},
                    {'type': 'text', 'content': '参考案例：西北窑洞元素与现代保温技术结合'}
                ]
            }
        },
        '适老化': {
            'score': 0,
            'threshold': 6,
            'suggestions': {
                'direction': '提升适老化设计细节',
                'cases': [
                    {'type': 'text', 'content': '参考案例：卫生间安装防滑扶手和紧急呼叫系统'},
                    {'type': 'text', 'content': '参考案例：厨房台面高度可调节设计，适应不同身高'},
                    {'type': 'text', 'content': '参考案例：地面无门槛设计，方便轮椅通行'}
                ]
            }
        },
        'AI工具运用': {
            'score': 0,
            'threshold': 6,
            'suggestions': {
                'direction': '提升 AI 工具在设计流程中的应用',
                'cases': [
                    {'type': 'text', 'content': '参考案例：使用 Midjourney 生成概念设计图'},
                    {'type': 'text', 'content': '参考案例：利用 AI 自动生成材料清单和预算表'},
                    {'type': 'text', 'content': '参考案例：运用 AI 分析户型并提供优化建议'}
                ]
            }
        }
    }
    
    # 模拟 AI 评分（实际应用中应调用大模型 API）
    import random
    for dimension in evaluation_dimensions:
        evaluation_dimensions[dimension]['score'] = random.randint(5, 10)
    
    # 生成 AI 评价文本
    ai_feedback_text = f"""【AI评价结果】

总体评分：⭐⭐⭐⭐
"""
    
    for dimension, data in evaluation_dimensions.items():
        stars = '⭐' * data['score']
        ai_feedback_text += f"""
{dimension}评分：{stars}（{data['score']}/10分）
"""
        if data['score'] < data['threshold']:
            ai_feedback_text += f"⚠️ 建议：{data['suggestions']['direction']}\n"
    
    ai_feedback_text += """
总体评价：这是一份优秀的作业设计，展现了良好的设计思维和专业能力。继续保持！
"""
    
    # 生成低分维度的修改清单
    modification_list = []
    for dimension, data in evaluation_dimensions.items():
        if data['score'] < data['threshold']:
            modification_list.append({
                'dimension': dimension,
                'score': data['score'],
                'direction': data['suggestions']['direction'],
                'cases': data['suggestions']['cases']
            })
    
    # 记录提交
    submission = {
        'filename': filename,
        'submit_time': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'file_path': file_path,
        'ai_feedback': ai_feedback_text,
        'evaluation_scores': evaluation_dimensions,
        'modification_list': modification_list
    }
    
    workflow_state['submissions'].append(submission)
    
    return jsonify({
        'success': True,
        'message': '作业提交成功！',
        'ai_feedback': ai_feedback_text,
        'evaluation_scores': evaluation_dimensions,
        'modification_list': modification_list
    })

@app.route('/generate_report', methods=['POST'])
def generate_report():
    """生成点评报告（测试版本）"""
    if not workflow_state['submissions']:
        return jsonify({'error': '还没有提交任何作业'}), 400
    
    # 生成模拟报告
    report = f"""
## AI作业点评报告

### 提交概览
- 总提交数：{len(workflow_state['submissions'])}
- 评估时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

### 详细评估
"""
    
    for i, submission in enumerate(workflow_state['submissions'], 1):
        report += f"""
#### 提交 {i}
- 文件名：{submission['filename']}
- 提交时间：{submission['submit_time']}
- 评估结果：✅ 提交成功，符合基本要求
- 改进建议：建议进一步优化方案设计，注重细节处理

{submission.get('ai_feedback', '暂无AI评价')}
"""
    
    workflow_state['feedback_report'] = report
    
    return jsonify({
        'success': True,
        'report': report
    })

@app.route('/reset_workflow', methods=['POST'])
def reset_workflow():
    """重置工作流"""
    workflow_state['stage'] = 2  # 修复：保持为2，允许用户直接进入学生端
    workflow_state['project_material'] = None
    workflow_state['course_content'] = None
    workflow_state['student_tasks'] = []
    workflow_state['submissions'] = []
    workflow_state['feedback_report'] = None
    
    return jsonify({
        'success': True,
        'message': '工作流已重置'
    })

if __name__ == '__main__':
    app.run(debug=False, host='0.0.0.0', port=5000)
